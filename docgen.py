import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import zipfile
from typing import Dict, List, Any
import openpyxl
from pathlib import Path
from docx import Document
from docx.shared import Inches
import tempfile
import os

# Page configuration
st.set_page_config(
    page_title="Neo Word Document Placeholder Processor",
    page_icon="📝",
    layout="wide"
)

class WordDocumentProcessor:
    def __init__(self):
        self.excel_data = {}      # (not used for multi‐doc, but kept for compatibility)
        self.template_doc = None
        self.placeholders = []
        
    def read_excel_data(self, uploaded_file) -> List[Dict[str, Any]]:
        try:
            # Read the entire sheet as strings, so blank cells become ""
            df_raw = pd.read_excel(uploaded_file, engine='openpyxl', dtype=str)
            df_raw = df_raw.fillna("")  # Replace NaN with empty string
            
            # We'll iterate row by row, grouping into blocks separated by fully blank rows
            blocks: List[List[tuple]] = []   # Each block is a list of (key, value) pairs
            current_block: List[tuple] = []
            
            # We assume the first two columns are the key/value columns
            # If the sheet has more than two columns, ignore columns beyond the first two
            for idx, row in df_raw.iterrows():
                key = str(row.iloc[0]).strip()
                val = str(row.iloc[1]).strip()
                
                # If both key and value are empty, treat as block separator
                if key == "" and val == "":
                    if current_block:
                        blocks.append(current_block)
                        current_block = []
                else:
                    # Only consider non‐empty key→value rows
                    if key != "" and val != "":
                        current_block.append((key, val))
                    # If key present but value empty, skip
                    elif key != "" and val == "":
                        continue
                    # If key empty but value present—unlikely—skip
                    else:
                        continue
            
            # After loop, if there's a leftover block, append it
            if current_block:
                blocks.append(current_block)
            
            # Now convert each block (list of (key,val)) into a dict, merging duplicate keys into lists
            data_dicts: List[Dict[str, Any]] = []
            for block in blocks:
                temp_dict: Dict[str, Any] = {}
                for (k, v) in block:
                    if k in temp_dict:
                        if isinstance(temp_dict[k], list):
                            temp_dict[k].append(v)
                        else:
                            temp_dict[k] = [temp_dict[k], v]
                    else:
                        temp_dict[k] = v
                data_dicts.append(temp_dict)
            
            return data_dicts
            
        except Exception as e:
            st.error(f"Error reading Excel file: {str(e)}")
            return []

    def extract_placeholders_from_doc(self, doc_file) -> List[str]:
        """Extract placeholders from Word document, in formats **<X>**, <X>, {X}, [X]."""
        try:
            doc = Document(doc_file)
            placeholders = set()
            
            # Patterns to search for
            patterns = [
                r'\*\*<([^>]+)>\*\*',   # **<placeholder>**
                r'<([^>]+)>',           # <placeholder>
                r'\{([^}]+)\}',         # {placeholder}
                r'\[([^\]]+)\]',        # [placeholder]
            ]
            
            # Paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                for pattern in patterns:
                    matches = re.findall(pattern, text)
                    placeholders.update(matches)
            
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        for pattern in patterns:
                            matches = re.findall(pattern, text)
                            placeholders.update(matches)
            
            return list(placeholders)
        
        except Exception as e:
            st.error(f"Error reading Word document: {str(e)}")
            return []
    
    def match_placeholders_to_data(
        self, placeholders: List[str], data_dict: Dict[str, Any]
    ) -> Dict[str, str]:
        """
        For each placeholder name (e.g. "Company Name", "Authorized Name"), find the best‐matching
        key in data_dict. Returns a mapping placeholder→data_key.
        """
        matches: Dict[str, str] = {}
        for placeholder in placeholders:
            p_clean = placeholder.strip()
            best_match = None
            best_score = 0.0
            
            # 1) Exact (case‐insensitive)
            for data_key in data_dict.keys():
                if p_clean.lower() == data_key.lower():
                    best_match = data_key
                    best_score = 1.0
                    break
            
            # 2) If no exact, measure simple similarity (# of overlapping words)
            if not best_match:
                for data_key in data_dict.keys():
                    score = self.calculate_similarity(p_clean, data_key)
                    if score > best_score and score > 0.6:
                        best_score = score
                        best_match = data_key
            
            if best_match:
                matches[p_clean] = best_match
        
        return matches
    
    def calculate_similarity(self, str1: str, str2: str) -> float:
        """
        Returns a float in [0,1] indicating simple similarity based on 
        exact match, substring, or word‐overlap ratio.
        """
        a = str1.lower().strip()
        b = str2.lower().strip()
        
        # Exact
        if a == b:
            return 1.0
        # Substring
        if a in b or b in a:
            return 0.9
        
        # Word‐overlap
        w1 = set(a.split())
        w2 = set(b.split())
        if w1 and w2:
            inter = w1.intersection(w2)
            union = w1.union(w2)
            return (len(inter) / len(union)) * 0.8
        
        return 0.0
    
    def dynamically_expand_word_template(
        self, doc_file, data_dict: Dict[str, Any]
    ) -> Document:
        """
        Given a Word template (with table rows containing placeholders like <Authorized Name>),
        auto‐detect how many "Authorized Name" or "Designation" entries are in data_dict (by
        checking if data_dict["Authorized Name"] is a list) and add extra table rows as needed.
        """
        try:
            doc_file.seek(0)
            doc = Document(doc_file)
            
            # 1) Build a map of base_field → how many entries we actually have
            field_counts: Dict[str, int] = {}
            for k, v in data_dict.items():
                low = k.lower()
                if 'authorized name' in low or 'designation' in low:
                    if isinstance(v, list):
                        base = k
                        cnt = len(v)
                    elif k.lower().endswith(tuple(f" {i}" for i in range(2,10))):
                        parts = k.rsplit(' ', 1)
                        base = parts[0]
                        try:
                            cnt = int(parts[-1])
                        except:
                            cnt = 1
                    else:
                        base = k
                        cnt = 1
                    
                    if base not in field_counts:
                        field_counts[base] = 0
                    field_counts[base] = max(field_counts[base], cnt)
            
            # 2) For every table in the doc, call auto_number_and_expand_table(...)
            for tbl in doc.tables:
                self.auto_number_and_expand_table(tbl, field_counts)
            
            return doc
        
        except Exception as e:
            st.error(f"Error expanding Word template: {str(e)}")
            return None
    
    def auto_number_and_expand_table(
        self, table, field_counts: Dict[str, int]
    ):
        """
        Within a single table, look for cells containing placeholders like <Authorized Name>, <Designation>.
        Auto‐number duplicates and then add extra rows if field_counts[...] > existing placeholder rows.
        """
        # (A) Collect all placeholder cells
        placeholder_cells = []
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if '<' in txt and '>' in txt:
                    ph = txt.replace('<', '').replace('>', '').strip()
                    placeholder_cells.append({
                        'row': r_idx,
                        'cell': c_idx,
                        'cell_obj': cell,
                        'placeholder': ph,
                        'orig_text': txt
                    })

        # (B) Auto‐number duplicates in the existing placeholders
        placeholder_counts: Dict[str, int] = {}
        for pc in placeholder_cells:
            name = pc['placeholder']
            placeholder_counts[name] = placeholder_counts.get(name, 0) + 1
            count = placeholder_counts[name]
            
            if count == 1:
                new_txt = f"<{name}>"
            else:
                new_txt = f"<{name} {count}>"
            
            pc['cell_obj'].text = new_txt
        
        # (C) Determine how many rows currently exist for Authorized Name / Designation
        existing_rows = max(
            placeholder_counts.get('Authorized Name', 1),
            placeholder_counts.get('Designation', 1)
        )
        
        # (D) Determine how many rows we need (from field_counts)
        max_needed = 1
        for base, cnt in field_counts.items():
            # If this table has that base placeholder at all
            if any(base.lower() in p.lower() for p in placeholder_counts.keys()):
                max_needed = max(max_needed, cnt)
        
        # (E) If we need more rows, append them
        if max_needed > existing_rows:
            to_add = max_needed - existing_rows
            
            # Find a "template row" (the last row containing any placeholder)
            template_row = None
            for ridx in reversed(range(len(table.rows))):
                row = table.rows[ridx]
                if any('<' in cell.text and '>' in cell.text for cell in row.cells):
                    template_row = row
                    break
            
            if template_row is not None:
                for i in range(to_add):
                    new_index = existing_rows + i + 1  # 1‐based numbering
                    new_row = table.add_row()
                    for c_idx, tmpl_cell in enumerate(template_row.cells):
                        if c_idx >= len(new_row.cells):
                            continue
                        new_cell = new_row.cells[c_idx]
                        ttxt = tmpl_cell.text.strip().lower()
                        
                        # If the template cell was <Authorized Name>, produce <Authorized Name N>
                        if '<authorized name>' in ttxt:
                            new_cell.text = f"<Authorized Name {new_index}>"
                        elif '<designation>' in ttxt:
                            new_cell.text = f"<Designation {new_index}>"
                        elif '<mop>' in ttxt:
                            new_cell.text = "<MOP>"
                        elif tmpl_cell.text.strip().isdigit():
                            # If template had a numeric serial, increment it
                            new_cell.text = str(new_index)
                        else:
                            # Signature cell or blank cell
                            new_cell.text = ""
    
    def populate_word_document(
        self, doc_file, data_dict: Dict[str, Any], matches: Dict[str, str]
    ) -> bytes:
        """
        """
        try:
            # 1) Expand template
            doc = self.dynamically_expand_word_template(doc_file, data_dict)
            if doc is None:
                return None
            
            # 2) Gather all placeholders in the expanded doc (paragraphs + tables)
            expanded_placeholders = []
            patterns = [
                r'\*\*<([^>]+)>\*\*',  # **<X>**
                r'<([^>]+)>',          # <X>
                r'\{([^}]+)\}',        # {X}
                r'\[([^\]]+)\]',       # [X]
            ]
            for para in doc.paragraphs:
                txt = para.text
                for pat in patterns:
                    found = re.findall(pat, txt)
                    expanded_placeholders.extend(found)
            
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        txt = cell.text
                        for pat in patterns:
                            found = re.findall(pat, txt)
                            expanded_placeholders.extend(found)
            
            expanded_placeholders = list(set(expanded_placeholders))
            
            # 3) Build a replacement_map for placeholders that map to single string (non‐list) values
            replacement_map: Dict[str, Any] = {}
            for ph, dkey in matches.items():
                if dkey in data_dict and not isinstance(data_dict[dkey], list):
                    replacement_map[ph] = data_dict[dkey]
            
            # If any expanded placeholder exactly matches a key, and that key is non-list, add it
            for ph in expanded_placeholders:
                if ph not in replacement_map:
                    for dkey in data_dict.keys():
                        if ph.lower() == dkey.lower() and not isinstance(data_dict[dkey], list):
                            replacement_map[ph] = data_dict[dkey]
                            break
            
            # 4) Replace in all paragraphs
            for para in doc.paragraphs:
                self.replace_placeholders_in_text(para, replacement_map)
            
            # 5) Replace in all tables with special handling for "Authorized Name N" / "Designation N"
            for tbl in doc.tables:
                self.populate_expanded_table_with_data(tbl, replacement_map, data_dict)
            
            # 6) Set some document properties
            doc.core_properties.title = "Board Resolution"
            doc.core_properties.author = "Document Generator"
            doc.core_properties.subject = "Populated Board Resolution"
            
            # 7) Save to bytes and return
            out_buffer = io.BytesIO()
            doc.save(out_buffer)
            out_buffer.seek(0)
            return out_buffer.getvalue()
        
        except Exception as e:
            st.error(f"Error processing Word document: {str(e)}")
            return None
    
    def populate_expanded_table_with_data(
        self, table, replacement_map: Dict[str, Any], data_dict: Dict[str, Any]
    ):
        """
        """
        mop_value = data_dict.get('MOP', 'severally/jointly')
        
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if '<' in txt and '>' in txt:
                    ph = txt.replace('<', '').replace('>', '').strip()  # e.g. "Authorized Name 2"
                    replacement_value = ""
                    
                    # 1) Direct in replacement_map?
                    if ph in replacement_map:
                        replacement_value = replacement_map[ph]
                    
                    else:
                        parts = ph.split()
                        last = parts[-1]
                        # 2) If last token is a digit, treat it as an index
                        if last.isdigit():
                            idx = int(last) - 1
                            base = ' '.join(parts[:-1])
                            if base in data_dict:
                                val = data_dict[base]
                                if isinstance(val, list):
                                    if 0 <= idx < len(val):
                                        replacement_value = val[idx]
                                else:
                                    replacement_value = val
                        else:
                            # 3) If ph is in data_dict and is a list, take [0]
                            if ph in data_dict:
                                val = data_dict[ph]
                                if isinstance(val, list):
                                    replacement_value = val[0]
                                else:
                                    replacement_value = val
                            # 4) Handle MOP specifically
                            elif ph.upper() == 'MOP':
                                replacement_value = mop_value
                            # 5) As a last resort, fuzzy‐match
                            else:
                                for dkey, dval in data_dict.items():
                                    if self.calculate_similarity(ph, dkey) > 0.8:
                                        if isinstance(dval, list):
                                            replacement_value = dval[0]
                                        else:
                                            replacement_value = dval
                                        break
                    
                    # If we found something, set the cell
                    if replacement_value:
                        self.clear_and_set_cell_text(cell, str(replacement_value))
    
    def clear_and_set_cell_text(self, cell, text: str):
        """
        Completely clear all runs in the cell and write `text` in a new run.
        """
        try:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.clear()
            if cell.paragraphs:
                cell.paragraphs[0].add_run(text)
            else:
                cell.add_paragraph().add_run(text)
        except Exception:
            cell.text = text
    
    def replace_placeholders_in_text(self, paragraph, replacement_map: Dict[str, Any]):
        """
        Replace placeholders in a single paragraph, preserving formatting.
        Patterns: **<X>**, <X>, {X}, [X].
        """
        patterns = [
            (r'\*\*<([^>]+)>\*\*', '**<{}>**'),
            (r'<([^>]+)>', '<{}>'),
            (r'\{([^}]+)\}', '{{{}}}'),
            (r'\[([^\]]+)\]', '[{}]'),
        ]
        
        for pattern, _ in patterns:
            matches = list(re.finditer(pattern, paragraph.text))
            for match in reversed(matches):
                ph = match.group(1).strip()
                if ph in replacement_map:
                    val = replacement_map[ph]
                    if isinstance(val, list):
                        if 'name' in ph.lower():
                            repl_text = '\n'.join(val)
                        else:
                            repl_text = ', '.join(val)
                    else:
                        repl_text = str(val)
                    start, end = match.start(), match.end()
                    self.replace_text_in_runs(paragraph, start, end, repl_text)
    
    def replace_text_in_runs(self, paragraph, start_pos: int, end_pos: int, replacement_text: str):
        """
        Given a run‐based paragraph, replace the text between start_pos and end_pos (character offsets)
        with replacement_text, preserving as much formatting as possible.
        """
        full_text = paragraph.text
        current = 0
        start_run = None
        end_run = None
        start_run_pos = 0
        end_run_pos = 0
        
        # Find which runs contain the start and end
        for i, run in enumerate(paragraph.runs):
            rlen = len(run.text)
            if current <= start_pos < current + rlen:
                start_run = i
                start_run_pos = start_pos - current
            if current < end_pos <= current + rlen:
                end_run = i
                end_run_pos = end_pos - current
                break
            current += rlen
        
        if start_run is not None and end_run is not None:
            if start_run == end_run:
                run = paragraph.runs[start_run]
                run.text = run.text[:start_run_pos] + replacement_text + run.text[end_run_pos:]
            else:
                # Replace across multiple runs
                paragraph.runs[start_run].text = (
                    paragraph.runs[start_run].text[:start_run_pos] + replacement_text
                )
                for j in range(start_run + 1, end_run):
                    paragraph.runs[j].text = ""
                if end_run < len(paragraph.runs):
                    paragraph.runs[end_run].text = paragraph.runs[end_run].text[end_run_pos:]


# Initialize the processor
if 'processor' not in st.session_state:
    st.session_state.processor = WordDocumentProcessor()

# Main UI
st.title("📝 Word Document Placeholder Processor (Multiple Documents)")
st.markdown("**Specialized for Board Resolutions and Legal Documents**")
st.markdown("Upload a single Excel file containing multiple record‐blocks (separated by blank rows). "
            "Each block will generate its own populated Word document, all zipped for download.")

col1, col2 = st.columns(2)
with col1:
    st.subheader("📊 Upload Excel Data (Multiple Records)")
    excel_file = st.file_uploader(
        "Choose Excel file with data (two‐column, blocks separated by blank rows)",
        type=['xlsx', 'xls'],
        help="Excel file: Column A = Field name, Column B = Value. Blank rows separate each record."
    )
with col2:
    st.subheader("📄 Upload Word Template")
    template_file = st.file_uploader(
        "Choose Word template (.docx)",
        type=['docx'],
        help="Word document with placeholders like **<Company Name>**, <Authorized Name>, <Designation>, <MOP>, etc."
    )

if excel_file and template_file:
    # 1) Read and split Excel into multiple data_dicts
    with st.spinner("📖 Reading Excel and splitting into multiple records..."):
        data_dicts = st.session_state.processor.read_excel_data(excel_file)
    
    if not data_dicts:
        st.error("❌ Could not extract any data records from the Excel file.")
        st.stop()
    
    st.success(f"✅ Found {len(data_dicts)} record(s) in Excel. Each will produce one Word document.")
    
    # 2) Extract placeholders from the Word template (once)
    with st.spinner("🔍 Extracting placeholders from Word template..."):
        placeholders = st.session_state.processor.extract_placeholders_from_doc(template_file)
    
    if not placeholders:
        st.error("❌ No placeholders found in the Word template. "
                 "Make sure you have placeholders like <Company Name>, <Authorized Name>, <Designation>, <MOP>, etc.")
        st.stop()
    
    st.subheader("🔖 Detected Placeholders in Template")
    placeholder_cols = st.columns(min(4, len(placeholders)))
    for i, ph in enumerate(placeholders):
        with placeholder_cols[i % len(placeholder_cols)]:
            st.code(ph)
    
    # 3) Show a summary of each record's "Company Name" (if present)
    st.subheader("📑 Preview of Uploaded Records")
    preview_md = ""
    for idx, dd in enumerate(data_dicts, start=1):
        company = dd.get("Company Name", "<No Company Name field>")
        preview_md += f"**Record {idx}:** {company}  \n"
    st.markdown(preview_md)
    
    # 4) Generate all documents when user clicks button
    if st.button(f"🚀 Generate {len(data_dicts)} Document(s)", type="primary"):
        docs_bytes_list: List[tuple] = []  # List of (filename, bytes)
        errors = 0
        
        for idx, dd in enumerate(data_dicts, start=1):
            # 4a) Build placeholder→data_key mapping for this block
            matches = st.session_state.processor.match_placeholders_to_data(placeholders, dd)
            
            # 4b) Populate the Word document for this block
            template_file.seek(0)
            doc_bytes = st.session_state.processor.populate_word_document(template_file, dd, matches)
            
            if doc_bytes is None:
                errors += 1
                continue
            
            # 4c) Build a filename: try to use Company Name, else use index
            comp = dd.get("Company Name", "").strip()
            if comp:
                # sanitize company name for filename
                safe_comp = re.sub(r'[^\w\s-]', '', comp).replace(" ", "_")[:30]
                filename = f"{safe_comp}_Resolution_{idx}.docx"
            else:
                filename = f"Document_{idx}.docx"
            
            docs_bytes_list.append((filename, doc_bytes))
        
        if errors > 0:
            st.warning(f"⚠️ {errors} document(s) failed to generate out of {len(data_dicts)}.")
        
        if not docs_bytes_list:
            st.error("❌ None of the documents could be generated. Please check your data/template.")
            st.stop()
        
        # 5) If there's exactly one document, offer it directly; otherwise, zip them all
        if len(docs_bytes_list) == 1:
            fname, content = docs_bytes_list[0]
            st.success(f"✅ Generated 1 document: **{fname}**")
            st.download_button(
                label="📄 Download Word Document",
                data=content,
                file_name=fname,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            # Create an in-memory ZIP
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, mode="w") as zf:
                for (fname, content) in docs_bytes_list:
                    zf.writestr(fname, content)
            zip_buffer.seek(0)
            
            st.success(f"✅ Generated {len(docs_bytes_list)} documents and compressed into a ZIP.")
            st.download_button(
                label="📦 Download All as ZIP",
                data=zip_buffer.getvalue(),
                file_name="Populated_Documents.zip",
                mime="application/zip"
            )
    
else:
    st.subheader("💡 How to Use")
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        **📊 Excel File Format (Two‐Column, Multiple Records)**  
        ```
        """)
    with col2:
        st.markdown("""
        **📄 Word Template Placeholders**""")
    st.markdown("---")
    st.markdown("🎯 **Specialized for Legal Documents:** Board Resolutions • Contracts • Agreements • Official Documents")
